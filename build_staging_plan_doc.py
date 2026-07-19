from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path

OUT = Path('outputs/Renivet_Staging_Environment_Plan.docx')
OUT.parent.mkdir(exist_ok=True)

NAVY='0B2545'; BLUE='2E74B5'; DARK='1F4D78'; MUTED='596579'; PALE='E8EEF5'; LIGHT='F4F6F9'; GREEN='EAF4EE'; AMBER='FFF5DE'; RED='FBEAEC'; WHITE='FFFFFF'; GREY='D9E1EA'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_border(cell, color='D9E1EA', size='6'):
    tcPr = cell._tc.get_or_add_tcPr(); borders = tcPr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right'):
        tag = 'w:' + edge; element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag); borders.append(element)
        element.set(qn('w:val'),'single'); element.set(qn('w:sz'), size); element.set(qn('w:color'),color)

def cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_width(cell, width):
    cell.width=Inches(width)
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(int(width*1440))); tcW.set(qn('w:type'),'dxa')

def set_font(run, size=11, color='000000', bold=False, italic=False):
    run.font.name='Calibri'; run._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
    run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold; run.italic=italic

def add_text(p, text, size=11, color='000000', bold=False, italic=False):
    r=p.add_run(text); set_font(r,size,color,bold,italic); return r

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def make_table(doc, rows, widths, header=True):
    table=doc.add_table(rows=len(rows), cols=len(widths)); table.alignment=WD_TABLE_ALIGNMENT.LEFT; table.autofit=False
    table.style='Table Grid'
    for ridx,row in enumerate(rows):
        for cidx,text in enumerate(row):
            cell=table.cell(ridx,cidx); set_width(cell,widths[cidx]); cell_margins(cell); set_cell_border(cell)
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
            add_text(p,str(text),9.2, WHITE if ridx==0 and header else '1C2837', bold=(ridx==0 and header))
            if ridx==0 and header: set_cell_shading(cell,NAVY)
            elif ridx%2==0 and ridx>0: set_cell_shading(cell,LIGHT)
    if header: set_repeat_table_header(table.rows[0])
    return table

def add_box(doc, title, text, fill=LIGHT, accent=BLUE):
    table=doc.add_table(rows=1, cols=1); table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    cell=table.cell(0,0); set_width(cell,6.5); cell_margins(cell,150,180,150,180); set_cell_shading(cell,fill); set_cell_border(cell,accent,'10')
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(3); add_text(p,title.upper(),9,accent,True)
    p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.12; add_text(p,text,10,'1C2837')
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def add_diagram(doc, rows, widths, caption):
    table=doc.add_table(rows=len(rows),cols=len(widths)); table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,row in enumerate(rows):
        for j,text in enumerate(row):
            c=table.cell(i,j); set_width(c,widths[j]); cell_margins(c,110,80,110,80); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
            if text in ('→','↓','↘','↗','↑'):
                add_text(p,text,18,BLUE,True)
            elif text:
                set_cell_shading(c, PALE if i%2==0 else LIGHT); set_cell_border(c,BLUE,'8'); add_text(p,text,9.3,NAVY,True)
            else:
                add_text(p,'',9)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8); add_text(p,caption,8.5,MUTED,italic=True)

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.9); sec.right_margin=Inches(0.9); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,color,before,after in [('Heading 1',16,BLUE,15,7),('Heading 2',13,BLUE,11,5),('Heading 3',11,DARK,8,4)]:
    st=styles[name]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

# Header/footer
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; header.paragraph_format.space_after=Pt(0); add_text(header,'RENIVET MARKETPLACE  |  STAGING ENVIRONMENT PLAN',8,MUTED,True)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.paragraph_format.space_before=Pt(0); add_text(footer,'Internal planning document  |  July 2026  |  ',8,MUTED)
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

# cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(54); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
add_text(p,'TECHNICAL DELIVERY PLAN',10,BLUE,True)
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(8); add_text(p,'Renivet Marketplace',30,NAVY,True)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(22); add_text(p,'Professional staging environment setup, release controls, and operating model',15,MUTED)
add_box(doc,'Recommended decision','Use a managed Vercel staging deployment at staging.renivet.com, backed by separate PostgreSQL, Redis, authentication, payment, email, storage, and analytics resources. Do not operate a shared or manually maintained VM unless a provider integration requires it.',PALE,BLUE)
make_table(doc,[['Document owner','Engineering / Platform'],['Audience','Engineering, QA, Operations, and business owners'],['Release scope','Next.js 15 marketplace, Drizzle/PostgreSQL, Redis, Clerk, Razorpay, Resend, Shiprocket, UploadThing, PostHog'],['Decision required','Approve isolated staging infrastructure and protected release workflow']], [1.55,4.95], header=False)
doc.add_page_break()

# summary
doc.add_heading('1. Executive summary',1)
doc.add_paragraph('Staging is the final production-like environment used to validate a release before customers see it. It must behave like production while remaining unable to affect production data, real customer communications, payments, fulfillment, or reporting.')
add_box(doc,'Success definition','A commit that passes automated checks can be deployed to staging, tested using safe test data and integrations, approved by QA, and promoted to production as the same immutable release.',GREEN,'277A4B')
doc.add_heading('Objectives',2)
for t in ['Deploy every approved change predictably with an auditable release path.','Validate checkout, webhooks, admin workflows, uploads, jobs, and integrations before production.','Prevent staging from reading or modifying production data and services.','Provide actionable logs, monitoring, backup/restore capability, and a documented rollback.']:
    doc.add_paragraph(t,style='List Bullet')
doc.add_heading('Environment model',2)
make_table(doc,[['Environment','URL / trigger','Purpose','Data policy'],['Local','localhost','Developer work','Synthetic or local data'],['Preview','Each pull request','Fast visual and functional review','Test data only'],['Staging','staging.renivet.com','Release candidate validation','Sanitized production-like data'],['Production','Live domain','Customer traffic','Live customer data']], [1.05,1.45,1.75,2.25])

doc.add_page_break()
doc.add_heading('2. Target architecture',1)
doc.add_paragraph('The staging deployment is deliberately isolated at every stateful or externally connected boundary. Vercel provides hosting, TLS, CDN, preview deployments, and deployment rollback; it does not remove the need for separate application data and vendor credentials.')
add_diagram(doc,[['Developer / QA','→','GitHub: staging branch','→','Vercel: staging.renivet.com'],['','','↓','','↓'],['Preview deployments','→','CI quality gates','→','Staging application'],['','','','','↓'],['','','PostgreSQL + Redis','←','Isolated vendor services']], [1.35,.35,1.55,.35,2.9], 'Figure 1. Staging deployment and service boundary.')
doc.add_heading('Staging service isolation',2)
make_table(doc,[['Component','Staging requirement','Control'],['Application','Separate Vercel project and staging hostname','Deployment Protection; HTTPS; noindex'],['PostgreSQL','Dedicated database and least-privilege database user','Backups; migration rehearsal; sanitized refresh'],['Redis','Dedicated Redis/Upstash database','Separate URL and access token'],['Identity','Separate Clerk instance and Svix signing secret','Register staging webhook endpoint'],['Payments','Razorpay test-mode account and webhook secret','No live charges or refunds'],['Email','Resend staging sender and recipient allow-list','Internal inboxes only'],['Fulfillment','Sandbox credentials or hard disablement','Test-order allow-list'],['Files / analytics','Separate UploadThing and PostHog resources','No production asset or metric contamination']], [1.25,3.4,1.85])

doc.add_page_break()
doc.add_heading('3. Release pipeline and change control',1)
doc.add_paragraph('The same checks must protect staging and production. A release is promoted by merging an approved, tested commit; it is not rebuilt or manually copied between servers.')
add_diagram(doc,[['Feature branch','→','Preview URL','→','Pull request review'],['','','↓','','↓'],['Automated checks','→','staging branch','→','staging.renivet.com'],['','','↓','','↓'],['QA approval','→','main branch','→','Production release']], [1.25,.35,1.35,.35,2.65], 'Figure 2. Controlled promotion path. The production release uses the approved staging commit.')
doc.add_heading('Mandatory CI gates',2)
for t in ['Install from the locked dependency file, then run lint, TypeScript validation, tests, and next build.','Run Drizzle migration validation against a disposable database copy before staging application.','Deploy only if all required checks pass; record commit SHA, migration version, and deployment URL.','Run post-deploy smoke tests: health, database/Redis connectivity, login, cart, test payment, webhook signature, upload, and admin authorization.']:
    doc.add_paragraph(t,style='List Bullet')
doc.add_heading('Branch and access policy',2)
make_table(doc,[['Asset','Policy'],['staging branch','Protected; pull request and passing CI required; engineering review required.'],['main branch','Protected; QA approval and passing CI required; only release owners can merge.'],['Vercel staging project','Platform administrators manage secrets; engineers receive least-privilege deploy access.'],['Database / Redis','No shared production credentials; individual access audited and time-bound where possible.']], [1.5,5.0])

doc.add_page_break()
doc.add_heading('4. Implementation runbook',1)
steps=[
('1. Establish hosting and DNS','Create the Vercel project renivet-marketplace-staging. Connect the staging branch, add staging.renivet.com, enforce Deployment Protection, and confirm TLS. Keep pull-request preview deployments enabled.'),
('2. Provision isolated data services','Create a dedicated PostgreSQL database, Redis database, and storage resource. Create least-privilege service users. Turn on backups and document restore ownership.'),
('3. Configure safe vendors','Create separate Clerk, Razorpay test, Resend, UploadThing, PostHog, and eligible logistics sandbox credentials. Register only staging webhook URLs.'),
('4. Load secrets','Populate Vercel staging environment variables only. Generate fresh random secrets; do not clone production values. Restrict visibility and maintain a named secret owner.'),
('5. Apply schema and seed test records','Back up staging, apply forward-only Drizzle migrations, and seed controlled test users, sellers, products, coupons, and orders.'),
('6. Add runtime safety guards','Introduce APP_ENV=staging. Display a staging banner, label generated artifacts, set noindex, and block or allow-list irreversible external actions.'),
('7. Enable workflow automation','Add protected CI, deployment, migration, smoke-test, alerting, and rollback procedures. Exercise them with a non-critical release.'),
('8. Sign off operations','Document ownership, on-call contacts, release checklist, incident steps, data refresh rules, and a quarterly access/secret review.')]
for title,body in steps:
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(1); add_text(p,title,11,DARK,True)
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.22); p.paragraph_format.space_after=Pt(4); add_text(p,body,10.2,'1C2837')

doc.add_page_break()
doc.add_heading('5. Configuration and operational safeguards',1)
doc.add_heading('Required application changes',2)
make_table(doc,[['Change','Why it matters','Acceptance criterion'],['Add APP_ENV','NODE_ENV should remain production on Vercel; staging needs an explicit safety signal.','APP_ENV permits local, preview, staging, production and is validated.'],['Add CRON_SECRET validation','Several cron routes rely on direct environment access.','All cron routes reject missing/invalid authorization.'],['Staging banner + labels','Prevents staff mistaking staging for production.','Persistent banner; emails, exports, PDFs labeled STAGING.'],['Robots protection','Prevents accidental search indexing.','noindex, nofollow response/header and no production canonical URL.'],['Outbound guard','Prevents financial, fulfillment, privacy, and communication mistakes.','Test account/order allow-list or integration disabled.']], [1.45,2.75,2.3])
doc.add_heading('Secret inventory',2)
doc.add_paragraph('Create distinct staging values for DATABASE_URL, REDIS_URL, Clerk and Svix keys, JWT_SECRET_KEY, Razorpay keys and webhook secret, Resend configuration, UploadThing token, Shiprocket/Unicommerce credentials, PostHog keys, Google/Facebook integration values, and all cron secrets. Store no secrets in Git or shared documents.')
add_box(doc,'Critical rule','Never configure staging with production DATABASE_URL, REDIS_URL, payment keys, email sender, shipping credentials, upload token, or analytics project. A different hostname alone is not isolation.',RED,'9B1C1C')
doc.add_heading('Cron and webhook policy',2)
for t in ['Enable only safe, authenticated staging schedules initially. Finance reconciliation, payout, deletion, operational email, and fulfillment jobs must be disabled or restricted to test fixtures.','Use separate webhook registrations and signing secrets for Clerk, Razorpay, and shipping events. Capture request IDs, signature validation result, and retry outcome in logs.','Set alerts for deployment failures, failed migrations, repeated webhook failures, job failures, database availability, and error-rate spikes.']:
    doc.add_paragraph(t,style='List Bullet')

doc.add_page_break()
doc.add_heading('6. Data protection, rollback, and release approval',1)
doc.add_heading('Database refresh process',2)
add_diagram(doc,[['Production backup','→','Sanitize / anonymize','→','Staging validation'],['','','↓','','↓'],['Restricted operator','→','Restore staging DB','→','QA confirms safe dataset']], [1.45,.35,1.65,.35,2.15], 'Figure 3. Controlled staging data refresh. No direct production database connection is permitted.')
doc.add_paragraph('Sanitize names, emails, phones, addresses, authentication identifiers, payment references, shipment data, support notes, and any operational attachment before a staging refresh. Limit access to trained operators, record each refresh, and use a restore point before importing.')
doc.add_heading('Rollback procedure',2)
make_table(doc,[['Scenario','Immediate action','Recovery owner'],['Application defect','Redeploy the last known-good Vercel deployment; pause promotion.','Release owner'],['Migration issue','Stop deployment, restore staging backup or execute pre-tested forward repair.','Database owner'],['Credential exposure','Disable/rotate affected staging credential, audit access, redeploy.','Platform owner'],['Unsafe integration behavior','Disable vendor credential or feature flag; preserve logs and test data.','Integration owner']], [1.3,3.65,1.55])
doc.add_heading('Release approval checklist',2)
for t in ['All CI gates passed and deployment references the intended commit SHA.','Migration backup and rehearsal completed; staging schema is current.','Critical journeys passed: authentication, catalog, cart, checkout, payment webhook, refunds, admin roles, uploads, shipping test, and scheduled job authorization.','No production data/service connection exists; outbound communications and fulfillment safety guard verified.','Errors, logs, dashboards, alerts, and rollback are available to named owners.','QA and release owner approve promotion to main.']:
    doc.add_paragraph(t,style='List Bullet')

doc.add_heading('7. Ownership and next actions',1)
make_table(doc,[['Timeframe','Action','Owner'],['Day 1','Create Vercel staging project, staging branch, hostname, and protected access.','Platform / Engineering'],['Days 1-2','Provision isolated PostgreSQL, Redis, and vendor test resources.','Platform + integration owners'],['Days 2-3','Load secrets, migrate schema, create test data, add APP_ENV safeguards.','Engineering'],['Days 3-4','Register safe webhooks, configure CI/smoke tests, enable monitoring.','Engineering + QA'],['Day 5','Run a complete release rehearsal, rollback test, and stakeholder sign-off.','Engineering + QA + Operations']], [1.1,3.85,1.55])
add_box(doc,'Final recommendation','Proceed with the managed Vercel architecture and separate service instances. It is the lowest-operations, professional solution for this Next.js application and creates a dependable path from preview to staging to production.',PALE,BLUE)

doc.core_properties.title='Renivet Marketplace Staging Environment Plan'
doc.core_properties.subject='Professional staging environment architecture and implementation runbook'
doc.core_properties.author='Renivet Engineering'
doc.save(OUT)
print(OUT.resolve())
